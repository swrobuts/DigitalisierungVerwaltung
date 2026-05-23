"""HTML→PDF via weasyprint + Jinja-Template."""
from datetime import date, datetime
from pathlib import Path
from typing import Any
from jinja2 import Environment, FileSystemLoader, select_autoescape
from weasyprint import HTML
from pruefung.models import PruefungsErgebnis


def _format_euro(v: Any) -> str:
    """1234.5 → '1.234,50 €'. Robust gegen str/decimal/float."""
    if v is None:
        return "—"
    try:
        n = float(v)
    except (ValueError, TypeError):
        return str(v)
    return f"{n:,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")


def _format_date(v: Any) -> str:
    """ISO-Datum oder date → '22.05.2026'."""
    if v is None:
        return "—"
    if isinstance(v, datetime):
        return v.strftime("%d.%m.%Y")
    if isinstance(v, date):
        return v.strftime("%d.%m.%Y")
    s = str(v)[:10]
    try:
        return datetime.strptime(s, "%Y-%m-%d").strftime("%d.%m.%Y")
    except ValueError:
        return s


_env = Environment(
    loader=FileSystemLoader(Path(__file__).parent / "templates"),
    autoescape=select_autoescape(["html", "j2"]),
)
_env.globals["format_euro"] = _format_euro
_env.globals["format_date"] = _format_date


def render_protokoll_pdf(antrag: dict, ergebnis: PruefungsErgebnis) -> bytes:
    tpl = _env.get_template("protokoll.html.j2")
    html = tpl.render(
        antrag=antrag,
        ergebnis=ergebnis,
        geprueft_am=datetime.now().strftime("%d.%m.%Y %H:%M"),
        verstoesse=ergebnis.anzahl_verstoesse(),
        hinweise=ergebnis.anzahl_hinweise(),
        status=ergebnis.pruefungsstatus(),
    )
    return HTML(string=html).write_pdf()


def render_bescheid_pdf(
    *,
    bescheid_id: str,
    antrag: dict,
    entscheidung: str,
    bewilligte_summe_euro: float | None,
    befunde: list[dict],
    bearbeiter_kommentar: str | None,
    ausgestellt_von: str | None,
    ausgestellt_am: datetime,
    doctree_version: str | None,
    geprueft_gegen: list[str] | None = None,
) -> bytes:
    """Rendert einen Verwaltungsbescheid (bewilligt/abgelehnt/rueckfrage).

    Die Befunde stammen aus dem letzten Prüfprotokoll und werden bereits mit
    angereichertem AHP-Wortlaut + Subsumtion übergeben (Caller konsultiert
    den Doctree). geprueft_gegen ist die Liste aller AHP-Sections, gegen die
    geprüft wurde — für Transparenz im Bescheid.
    """
    tpl = _env.get_template("bescheid.html.j2")
    # Pfad zum Würzburg-Wappen (SVG liegt im templates/-Ordner).
    # WeasyPrint braucht ein absolutes file://-URI.
    wappen_path = (Path(__file__).parent / "templates" / "wuerzburg_coa.svg").resolve()
    html = tpl.render(
        bescheid_id=bescheid_id,
        antrag=antrag,
        entscheidung=entscheidung,
        bewilligte_summe_euro=bewilligte_summe_euro,
        befunde=befunde,
        bearbeiter_kommentar=bearbeiter_kommentar,
        ausgestellt_von=ausgestellt_von,
        ausgestellt_am=ausgestellt_am,
        doctree_version=doctree_version or "—",
        geprueft_gegen=geprueft_gegen or [],
        wappen_src=wappen_path.as_uri(),
    )
    # base_url muss gesetzt sein, damit WeasyPrint relative/file-URLs auflöst.
    return HTML(string=html, base_url=str(wappen_path.parent)).write_pdf()


# ── DOCX-Helfer ────────────────────────────────────────────────────────
# python-docx hat keine High-Level-API für Shading, Borders auf Zellen /
# Paragraphen oder Spaltenbreiten. Wir setzen das direkt im OXML — das ist
# stabil über alle Word/LibreOffice-Versionen und sieht aus wie das PDF.

# Würzburg-CI
_CI_ROT = "AD0E36"
_CI_DUNKEL = "1a1a1a"
_CI_GRAU = "555555"
_CI_HELLGRAU = "888888"
_CI_BG_GRAU = "FAFAFA"

_ENTSCHEIDUNG_FARBEN = {
    "bewilligt":  {"border": "047857", "bg": "ECFDF5", "text": "064E3B"},
    "abgelehnt":  {"border": "B91C1C", "bg": "FEF2F2", "text": "7F1D1D"},
    "rueckfrage": {"border": "B45309", "bg": "FFFBEB", "text": "78350F"},
}


def _docx_set_cell_shading(cell, hex_color: str) -> None:
    """Hintergrundfarbe einer Tabellenzelle setzen (kein public API in python-docx)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _docx_set_cell_borders(
    cell,
    *,
    color: str = "AAAAAA",
    size: int = 4,
    sides: tuple[str, ...] = ("top", "left", "bottom", "right"),
) -> None:
    """Rahmen einer Zelle einzeln pro Seite setzen. size in 1/8-Punkt
    (also size=8 ≈ 1pt). sides aus {top,left,bottom,right}."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _docx_paragraph_border(
    paragraph,
    *,
    side: str = "bottom",
    color: str = _CI_ROT,
    size: int = 12,
    space: int = 4,
) -> None:
    """Border auf einer Seite eines Absatzes — z.B. Trennlinie unter dem
    Briefkopf (side='bottom') oder farbiger Streifen links neben einem
    Begründungs-Item (side='left'). size in 1/8-Punkt."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    # pBdr muss vor anderen Children am richtigen Schema-Index stehen,
    # sonst akzeptiert Word es nicht. Wir hängen vorhandenen pBdr wieder
    # auf, statt blind anzuhängen.
    existing = pPr.find(qn("w:pBdr"))
    if existing is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.insert(0, pBdr)
    else:
        pBdr = existing
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    pBdr.append(border)


# Backwards-compat-Wrapper — Aufrufer aus älteren Edits.
def _docx_paragraph_border_bottom(
    paragraph, *, color: str = _CI_ROT, size: int = 12,
) -> None:
    _docx_paragraph_border(paragraph, side="bottom", color=color, size=size)


def _docx_paragraph_left_indent(paragraph, cm: float) -> None:
    """Setzt links-Einrückung — wird für 'Streifen-Item' kombiniert mit
    border-left, sodass der Text-Inhalt rechts neben der farbigen Linie steht."""
    from docx.shared import Cm
    paragraph.paragraph_format.left_indent = Cm(cm)


def _signature_name(value: str | None) -> str:
    """Aus 'vorname.nachname@thws.de' wird 'Vorname Nachname'. Verhindert
    dass dieselbe Email zweimal im Bescheid auftaucht (einmal im Briefkopf
    rechts als Meta, einmal unter der Unterschrift). Wenn der Wert keine
    Email ist, bleibt er unverändert."""
    if not value:
        return "Im Auftrag · Sachbearbeitung Sozialreferat"
    if "@" in value:
        local = value.split("@", 1)[0]
        parts = [p.capitalize() for p in local.replace("_", ".").split(".") if p]
        return " ".join(parts) if parts else value
    return value


def _docx_set_col_widths(table, widths_cm: list[float]) -> None:
    """Spaltenbreiten setzen — python-docx vergisst das gelegentlich, also
    pro Zelle nochmal explizit drüber."""
    from docx.shared import Cm
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def render_bescheid_docx(
    *,
    bescheid_id: str,
    antrag: dict,
    entscheidung: str,
    bewilligte_summe_euro: float | None,
    befunde: list[dict],
    bearbeiter_kommentar: str | None,
    ausgestellt_von: str | None,
    ausgestellt_am: datetime,
    doctree_version: str | None,
    geprueft_gegen: list[str] | None = None,
) -> bytes:
    """Editierbare DOCX-Variante mit demselben Layout wie das PDF.

    Unterschiede zur PDF nur dort wo Word es zwingend verlangt: Stempel-
    Platzhalter wird weggelassen (Word setzt keinen gestrichelten Rahmen
    automatisch um leere Zellen), und einige Mikro-Abstände sind in mm
    statt pt. Der Sachbearbeiter kann das DOCX direkt vor dem Versand in
    Word/LibreOffice noch anpassen (Briefkopf-Logo tauschen, Anschrift
    korrigieren, Norm-Zitat ergänzen)."""
    import io
    from pathlib import Path
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Mm, Pt, RGBColor

    doc = Document()

    # ── Page-Setup (DIN-5008 simuliert) ────────────────────────────
    for section in doc.sections:
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(28)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)

    # ── Default-Font ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── 1. BRIEFKOPF: Wappen | Stadt-Block | Meta ──────────────────
    # 3-spaltige Tabelle ohne Borders, mit Wappen links, Stadt-Block
    # mittig, Meta-Daten rechts. Darunter eine rote 1.5pt-Linie.
    kopf = doc.add_table(rows=1, cols=3)
    kopf.autofit = False
    _docx_set_col_widths(kopf, [2.5, 11.0, 5.0])
    kopf_l, kopf_m, kopf_r = kopf.rows[0].cells

    # Wappen einbetten (PNG-Variante des SVG — python-docx kann kein SVG).
    wappen_png = Path(__file__).parent / "templates" / "wuerzburg_coa.png"
    if wappen_png.exists():
        wappen_p = kopf_l.paragraphs[0]
        wappen_p.add_run().add_picture(str(wappen_png), width=Cm(2.0))

    # Stadt-Block (CI-Rot Stadt, dunkelgrau Ressort + Adresse)
    stadt_p = kopf_m.paragraphs[0]
    stadt_r = stadt_p.add_run("STADT WÜRZBURG")
    stadt_r.bold = True
    stadt_r.font.size = Pt(14)
    stadt_r.font.color.rgb = RGBColor.from_string(_CI_ROT)
    ressort_p = kopf_m.add_paragraph()
    ressort_r = ressort_p.add_run(
        "Sozialreferat — Fachbereich Integration, Inklusion und Senioren"
    )
    ressort_r.font.size = Pt(9.5)
    ressort_r.font.color.rgb = RGBColor.from_string(_CI_DUNKEL)
    adresse_p = kopf_m.add_paragraph()
    adresse_r = adresse_p.add_run(
        "Karmelitenstraße 43 · 97070 Würzburg · sozialreferat@stadt.wuerzburg.de"
    )
    adresse_r.font.size = Pt(8.5)
    adresse_r.font.color.rgb = RGBColor.from_string(_CI_GRAU)

    # Meta-Daten rechts (AZ, Datum, SB)
    def _add_meta(cell, label: str, value: str, monospace: bool = False) -> None:
        lbl_p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
        lbl_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lbl_r = lbl_p.add_run(label.upper())
        lbl_r.font.size = Pt(7.5)
        lbl_r.font.color.rgb = RGBColor.from_string(_CI_HELLGRAU)
        val_p = cell.add_paragraph()
        val_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        val_r = val_p.add_run(value)
        val_r.font.size = Pt(9.5)
        val_r.font.color.rgb = RGBColor.from_string(_CI_DUNKEL)
        if monospace:
            val_r.font.name = "Consolas"
            val_r.bold = True

    _add_meta(kopf_r, "Aktenzeichen", antrag.get("antragsnummer", "—"), monospace=True)
    _add_meta(kopf_r, "Datum", _format_date(ausgestellt_am))
    if ausgestellt_von:
        _add_meta(kopf_r, "Sachbearbeitung", _signature_name(ausgestellt_von))

    # Rote Trennlinie unter dem Briefkopf
    trenn = doc.add_paragraph()
    trenn.paragraph_format.space_before = Pt(0)
    trenn.paragraph_format.space_after = Pt(4)
    _docx_paragraph_border_bottom(trenn, color=_CI_ROT, size=16)

    # ── 2. ANSCHRIFTENFELD (DIN-5008-Position B) ───────────────────
    absender_p = doc.add_paragraph()
    absender_p.paragraph_format.space_before = Pt(14)
    absender_r = absender_p.add_run(
        "Stadt Würzburg · Sozialreferat · Karmelitenstraße 43 · 97070 Würzburg"
    )
    absender_r.font.size = Pt(7.5)
    absender_r.font.color.rgb = RGBColor.from_string(_CI_HELLGRAU)
    _docx_paragraph_border_bottom(absender_p, color="AAAAAA", size=4)

    emp_p = doc.add_paragraph()
    emp_p.paragraph_format.space_before = Pt(3)
    # Engerer Zeilenabstand für den Empfänger-Block — sieht im Brief
    # eines Verwaltungsamts professioneller aus als der lockere
    # Default-Abstand.
    emp_p.paragraph_format.line_spacing = 1.15
    name_r = emp_p.add_run(antrag.get("traeger", ""))
    name_r.bold = True
    name_r.font.size = Pt(11)
    if antrag.get("ansprechpartner"):
        emp_p.add_run(f"\n{antrag['ansprechpartner']}").font.size = Pt(11)
    emp_p.add_run(
        f"\n{antrag.get('strasse','')} {antrag.get('hausnummer','')}"
    ).font.size = Pt(11)
    emp_p.add_run(
        f"\n{antrag.get('plz','')} {antrag.get('ort','')}"
    ).font.size = Pt(11)

    # ── 3. BETREFF ─────────────────────────────────────────────────
    betreff_titel = {
        "bewilligt":  f"Bewilligung Ihres Förderantrags vom {_format_date(antrag.get('antragsdatum'))}",
        "abgelehnt":  f"Ablehnung Ihres Förderantrags vom {_format_date(antrag.get('antragsdatum'))}",
        "rueckfrage": f"Rückfrage zu Ihrem Förderantrag vom {_format_date(antrag.get('antragsdatum'))}",
    }.get(entscheidung, "Bescheid")

    betreff_p = doc.add_paragraph()
    betreff_p.paragraph_format.space_before = Pt(14)
    btr_r = betreff_p.add_run(betreff_titel)
    btr_r.bold = True
    btr_r.font.size = Pt(11)

    sub_p = doc.add_paragraph()
    sub_r = sub_p.add_run(
        f"Altentagesstätten — APL 2 · Haushaltsjahr {antrag.get('haushaltsjahr','—')} · "
        f"Rechtsgrundlage: AHP-Förderrichtlinie der Stadt Würzburg "
        f"(Beschluss vom 27.03.2025)"
    )
    sub_r.font.size = Pt(9.5)
    sub_r.font.color.rgb = RGBColor.from_string(_CI_GRAU)

    # ── 4. ANREDE + ERÖFFNUNGSSATZ ─────────────────────────────────
    anr_p = doc.add_paragraph("Sehr geehrte Damen und Herren,")
    anr_p.paragraph_format.space_before = Pt(10)

    eroeff = doc.add_paragraph()
    eroeff.add_run("zu Ihrem Antrag auf Förderung der Einrichtung ")
    eroeff.add_run(antrag.get("name", "")).bold = True
    eroeff.add_run(
        " ergeht nach Prüfung gemäß der AHP-Förderrichtlinie der Stadt "
        "Würzburg folgende Entscheidung:"
    )

    # ── 5. ENTSCHEIDUNGS-BLOCK (farbig gerahmt + Shading) ──────────
    farben = _ENTSCHEIDUNG_FARBEN.get(entscheidung, _ENTSCHEIDUNG_FARBEN["rueckfrage"])
    ent_tbl = doc.add_table(rows=1, cols=1)
    ent_tbl.autofit = False
    _docx_set_col_widths(ent_tbl, [16.5])
    ent_cell = ent_tbl.rows[0].cells[0]
    _docx_set_cell_shading(ent_cell, farben["bg"])
    _docx_set_cell_borders(ent_cell, color=farben["border"], size=12)
    # Innen-Padding via Margins
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = ent_cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), "180")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    lbl_p = ent_cell.paragraphs[0]
    lbl_r = lbl_p.add_run("ENTSCHEIDUNG")
    lbl_r.font.size = Pt(8.5)
    lbl_r.font.color.rgb = RGBColor.from_string(_CI_GRAU)
    verb_p = ent_cell.add_paragraph()
    verb_text = {
        "bewilligt":  "Antrag wird BEWILLIGT.",
        "abgelehnt":  "Antrag wird ABGELEHNT.",
        "rueckfrage": "Vor abschließender Entscheidung sind RÜCKFRAGEN zu klären.",
    }.get(entscheidung, "")
    verb_r = verb_p.add_run(verb_text)
    verb_r.bold = True
    verb_r.font.size = Pt(14)
    verb_r.font.color.rgb = RGBColor.from_string(farben["text"])
    if entscheidung == "bewilligt" and bewilligte_summe_euro:
        summe_p = ent_cell.add_paragraph()
        summe_r = summe_p.add_run(
            f"Bewilligte Fördersumme: {_format_euro(bewilligte_summe_euro)}"
        )
        summe_r.bold = True
        summe_r.font.size = Pt(12)
        summe_r.font.color.rgb = RGBColor.from_string(farben["text"])

    # ── 6. I. SACHVERHALT (nur wenn nicht bewilligt) ───────────────
    def _h2(text: str) -> None:
        """h2-äquivalent: 11pt fett mit dünner grauer Linie darunter."""
        h_p = doc.add_paragraph()
        h_p.paragraph_format.space_before = Pt(14)
        h_p.paragraph_format.space_after = Pt(4)
        h_r = h_p.add_run(text)
        h_r.bold = True
        h_r.font.size = Pt(11)
        h_r.font.color.rgb = RGBColor.from_string(_CI_DUNKEL)
        _docx_paragraph_border_bottom(h_p, color="AAAAAA", size=4)

    if entscheidung != "bewilligt":
        _h2("I. Sachverhalt")
        doc.add_paragraph("Folgende Antragsdaten lagen der Prüfung zugrunde:")

        kvs = [
            ("Aktenzeichen", antrag.get("antragsnummer", "—"), True),
            ("Antragsdatum", _format_date(antrag.get("antragsdatum")), False),
            ("Haushaltsjahr", str(antrag.get("haushaltsjahr", "—")), False),
            ("Träger", antrag.get("traeger", "—"), False),
            ("Einrichtung", antrag.get("name", "—"), False),
            ("Sitz",
             f"{antrag.get('strasse','')} {antrag.get('hausnummer','')}, "
             f"{antrag.get('plz','')} {antrag.get('ort','')}", False),
            ("Bankverbindung",
             f"{antrag.get('bankverbindung','—')} · IBAN {antrag.get('iban','—')}",
             False),
        ]
        if antrag.get("foerderbereich"):
            kvs.append(("Förderbereich (laut Antrag)", str(antrag["foerderbereich"]), False))
        if antrag.get("geforderte_foerdersumme_euro") is not None:
            kvs.append((
                "Geforderte Fördersumme",
                f"{_format_euro(antrag['geforderte_foerdersumme_euro'])} / Jahr",
                False,
            ))
        if antrag.get("stadtbewohner_anteil") is not None:
            anteil = int(round(float(antrag["stadtbewohner_anteil"]) * 100))
            kvs.append(("Anteil Würzburger Teilnehmer", f"{anteil} %", False))

        sv_tbl = doc.add_table(rows=len(kvs), cols=2)
        sv_tbl.autofit = False
        _docx_set_col_widths(sv_tbl, [5.0, 11.5])
        for i, (label, value, monospace) in enumerate(kvs):
            l_cell, v_cell = sv_tbl.rows[i].cells
            _docx_set_cell_borders(l_cell, color="EEEEEE", size=2, sides=("bottom",))
            _docx_set_cell_borders(v_cell, color="EEEEEE", size=2, sides=("bottom",))
            l_p = l_cell.paragraphs[0]
            l_r = l_p.add_run(label.upper())
            l_r.font.size = Pt(8.5)
            l_r.font.color.rgb = RGBColor.from_string(_CI_GRAU)
            v_p = v_cell.paragraphs[0]
            v_r = v_p.add_run(str(value))
            v_r.font.size = Pt(10)
            if monospace:
                v_r.font.name = "Consolas"

        _h2("II. Durchgeführte Prüfung")
        doc.add_paragraph(
            f"Der Antrag wurde am {_format_date(ausgestellt_am)} durch das "
            f"Sozialreferat einer dreistufigen Prüfung gegen die AHP-Förder"
            f"richtlinie der Stadt Würzburg (Beschluss vom 27.03.2025, "
            f"Doctree-Version {doctree_version or '—'}) unterzogen:"
        )
        for nr, schritt in [
            ("1.", "Formal-strukturelle Prüfung — Vollständigkeit + technische "
                   "Validität (IBAN-Prüfziffer nach ISO 13616, Postleitzahl-"
                   "Format, E-Mail-Format, Plausibilität Haushaltsjahr)."),
            ("2.", "Inhaltliche Konformitätsprüfung (Ontologie) — Abgleich aller "
                   "AHP-belegten Cap-Regeln, Fristen, Anteilskalkulationen und "
                   "Pflichtangaben gegen die Antragswerte."),
            ("3.", "Wortlaut-basierte AHP-Prüfung — KI-gestützter Abgleich gegen "
                   "den Volltext der AHP-Richtlinie (Doctree-Navigation + "
                   "semantische Suche)."),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.add_run(f"{nr} ").bold = True
            p.add_run(schritt)

        if geprueft_gegen:
            gg_p = doc.add_paragraph()
            gg_p.paragraph_format.space_before = Pt(4)
            gg_r = gg_p.add_run(
                "Geprüft wurde gegen folgende AHP-Stellen: "
                + ", ".join(geprueft_gegen)
            )
            gg_r.font.size = Pt(9.5)
            gg_r.font.color.rgb = RGBColor.from_string(_CI_GRAU)

    # ── 7. III. RECHTLICHE WÜRDIGUNG — Verstöße mit Subsumtion ────
    verstoesse = [b for b in befunde if b.get("schwere") == "verstoss"]
    if verstoesse:
        prefix = "III. " if entscheidung != "bewilligt" else ""
        _h2(f"{prefix}Rechtliche Würdigung — festgestellte Verstöße")
        if entscheidung == "abgelehnt":
            plural = "Verstoß" if len(verstoesse) == 1 else f"{len(verstoesse)} Verstöße"
            doc.add_paragraph(
                f"Die Ablehnung stützt sich auf {plural} gegen Bestimmungen der "
                f"AHP-Förderrichtlinie. Jeder Befund wird im Folgenden einzeln "
                f"dargelegt — bestehend aus dem konkreten Sachverhalt, dem "
                f"wörtlichen Auszug der einschlägigen Norm, sowie der "
                f"rechtlichen Würdigung:"
            )
        else:
            doc.add_paragraph("Vor abschließender Entscheidung sind folgende Punkte zu klären:")

        for i, b in enumerate(verstoesse, start=1):
            # 2-spaltige Tabelle: schmale rote Streifen-Zelle | Inhalt.
            # Ergibt im DOCX denselben Effekt wie border-left:2pt im PDF.
            item = doc.add_table(rows=1, cols=2)
            item.autofit = False
            _docx_set_col_widths(item, [0.15, 16.35])
            stripe, content = item.rows[0].cells
            _docx_set_cell_shading(stripe, "B91C1C")
            # Border-Container sorgt für scharfe Kante
            _docx_set_cell_borders(stripe, color="B91C1C", size=2)

            titel_p = content.paragraphs[0]
            nr_r = titel_p.add_run(f"{i}.  ")
            nr_r.bold = True
            nr_r.font.size = Pt(11)
            nr_r.font.color.rgb = RGBColor.from_string("B91C1C")
            t_r = titel_p.add_run(b.get("beschreibung", ""))
            t_r.bold = True
            t_r.font.size = Pt(11)

            def _feld(label: str, body: str, italic: bool = False, framed: bool = False) -> None:
                lbl_p = content.add_paragraph()
                lbl_p.paragraph_format.space_before = Pt(3)
                lbl_r = lbl_p.add_run(label.upper())
                lbl_r.font.size = Pt(7.5)
                lbl_r.font.color.rgb = RGBColor.from_string("777777")
                lbl_r.bold = True
                body_p = content.add_paragraph()
                body_r = body_p.add_run(body)
                body_r.font.size = Pt(10)
                if framed:
                    # eingerahmtes Norm-Zitat: italic + linke graue Border
                    # + Einrückung (alles via paragraph-properties, ohne
                    # Sub-Tabelle — die würde im DOCX-Modell aus der
                    # umschließenden Zelle herausspringen, da add_table
                    # immer an doc.body anhängt).
                    body_r.italic = True
                    body_r.font.color.rgb = RGBColor.from_string("444444")
                    _docx_paragraph_border(body_p, side="left", color="888888", size=8, space=8)
                    _docx_paragraph_left_indent(body_p, 0.3)
                elif italic:
                    body_r.italic = True

            if b.get("sachverhalt"):
                _feld("Sachverhalt im Einzelnen", b["sachverhalt"])
            if b.get("ahp_wortlaut"):
                label = "Rechtsgrundlage · Wortlaut"
                if b.get("paragraph_ref"):
                    label += f" · {b['paragraph_ref']}"
                _feld(label, b["ahp_wortlaut"], framed=True)
            if b.get("wuerdigung"):
                _feld("Rechtliche Würdigung", b["wuerdigung"])

            # Mini-Abstand zwischen den Items
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)

    # ── 8. IV. KONSEQUENZ / AUFLAGEN ───────────────────────────────
    if entscheidung == "abgelehnt":
        _h2("IV. Entscheidungsgründe und Konsequenz")
        plural = "Verstoßes" if len(verstoesse) == 1 else "Verstöße"
        doc.add_paragraph(
            f"Auf Grund der unter Ziffer III dargestellten {len(verstoesse)} "
            f"{plural} gegen die AHP-Förderrichtlinie ist der Antrag nach "
            f"pflichtgemäßer Prüfung abzulehnen. Eine teilweise Bewilligung "
            f"kommt nicht in Betracht, da die festgestellten Verstöße "
            f"entweder formaler Natur sind (technische Voraussetzungen wie "
            f"IBAN-Validität, Antragsfrist, Trägersitz) oder die beantragte "
            f"Förderhöhe materiell überschreiten."
        )
        doc.add_paragraph(
            "Im Übrigen besteht nach AHP Kap. 3.4 grundsätzlich kein "
            "Rechtsanspruch auf eine Förderung; über Art und Höhe der "
            "Bewilligung entscheidet der Sozialausschuss der Stadt Würzburg "
            "nach Prüfung durch die Leitung Seniorenarbeit (Fachbereich IIS)."
        )

    if entscheidung == "bewilligt":
        _h2("Auflagen zur Mittelverwendung")
        try:
            hj_next = int(antrag.get("haushaltsjahr", 0)) + 1
        except (TypeError, ValueError):
            hj_next = "—"
        auflagen = [
            ("Mittelverwendung", "Die bewilligten Mittel sind ausschließlich "
                "zweckgebunden für die geförderte Maßnahme zu verwenden "
                "(AHP Kap. 3.8)."),
            ("Verwendungsnachweis", f"Bis zum 1. April {hj_next} ist ein "
                "vollständiger Verwendungsnachweis (sachlicher Bericht + "
                "zahlenmäßiger Nachweis) einzureichen (AHP Kap. 3.8)."),
            ("Belege", "Originalbelege sind mindestens 3 Jahre nach Vorlage "
                "des Verwendungsnachweises aufzubewahren (AHP Kap. 3.8)."),
            ("Logo-Pflicht", "Bei Veröffentlichungen ist das Logo der Stadt "
                "Würzburg zu verwenden und der Erhalt von Fördergeldern nach "
                "SGB XII § 71 zu vermerken (AHP Kap. 2)."),
            ("Prüfungsrecht", "Die Stadt Würzburg ist berechtigt, die "
                "ordnungsgemäße Verwendung zu prüfen. Sie sind verpflichtet, "
                "Einsicht in alle Bücher und Belege zu gewähren (AHP Kap. 3.9)."),
        ]
        for nr, (label, body) in enumerate(auflagen, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3)
            nr_r = p.add_run(f"{nr}. ")
            nr_r.bold = True
            lbl_r = p.add_run(f"{label}: ")
            lbl_r.bold = True
            p.add_run(body)

    if entscheidung == "rueckfrage":
        rf_p = doc.add_paragraph(
            "Bitte ergänzen Sie die fehlenden Angaben innerhalb von vier "
            "Wochen ab Zugang dieses Schreibens. Nach Eingang der Ergänzungen "
            "wird die Prüfung fortgesetzt."
        )
        rf_p.paragraph_format.space_before = Pt(8)

    # ── 9. SACHBEARBEITER-KOMMENTAR ────────────────────────────────
    if bearbeiter_kommentar:
        komm_tbl = doc.add_table(rows=1, cols=1)
        komm_tbl.autofit = False
        _docx_set_col_widths(komm_tbl, [16.5])
        k_cell = komm_tbl.rows[0].cells[0]
        _docx_set_cell_shading(k_cell, "F7F6F4")
        _docx_set_cell_borders(k_cell, color=_CI_ROT, size=4, sides=("left",))
        lbl_p = k_cell.paragraphs[0]
        lbl_r = lbl_p.add_run("ANMERKUNG DER SACHBEARBEITUNG")
        lbl_r.bold = True
        lbl_r.font.size = Pt(7.5)
        lbl_r.font.color.rgb = RGBColor.from_string(_CI_HELLGRAU)
        body_p = k_cell.add_paragraph()
        body_r = body_p.add_run(bearbeiter_kommentar)
        body_r.font.size = Pt(10)

    # ── 10. SCHLUSSFORMEL + UNTERSCHRIFTSLINIE ─────────────────────
    schluss_p = doc.add_paragraph("Mit freundlichen Grüßen")
    schluss_p.paragraph_format.space_before = Pt(14)

    # Zwei Leerzeilen als Unterschriftraum (statt vorher drei — sah zu
    # luftig aus). Sachbearbeiter kann in Word zusätzliche Returns
    # einfügen, falls er einen größeren Stempelraum braucht.
    for _ in range(2):
        doc.add_paragraph()

    sig_p = doc.add_paragraph()
    sig_r = sig_p.add_run(_signature_name(ausgestellt_von))
    sig_r.font.size = Pt(9.5)
    sig_r.font.color.rgb = RGBColor.from_string(_CI_DUNKEL)
    # Trennlinie ÜBER dem Namen → simuliert Unterschriftslinie
    _docx_paragraph_border(sig_p, side="top", color=_CI_DUNKEL, size=4)

    # ── 11. RECHTSBEHELFSBELEHRUNG (graue Box) ─────────────────────
    rb_tbl = doc.add_table(rows=1, cols=1)
    rb_tbl.autofit = False
    _docx_set_col_widths(rb_tbl, [16.5])
    rb_cell = rb_tbl.rows[0].cells[0]
    _docx_set_cell_shading(rb_cell, _CI_BG_GRAU)
    _docx_set_cell_borders(rb_cell, color="888888", size=4)

    rb_h_p = rb_cell.paragraphs[0]
    rb_h_r = rb_h_p.add_run("RECHTSBEHELFSBELEHRUNG")
    rb_h_r.bold = True
    rb_h_r.font.size = Pt(10)
    rb_h_r.font.color.rgb = RGBColor.from_string(_CI_DUNKEL)
    rb_p1 = rb_cell.add_paragraph()
    rb_p1_r = rb_p1.add_run(
        "Gegen diesen Bescheid kann innerhalb eines Monats nach Bekanntgabe "
        "Widerspruch erhoben werden. Der Widerspruch ist schriftlich oder "
        "zur Niederschrift bei der Stadt Würzburg, Sozialreferat, "
        "Karmelitenstraße 43, 97070 Würzburg einzulegen."
    )
    rb_p1_r.font.size = Pt(9)
    rb_p2 = rb_cell.add_paragraph()
    rb_p2_lbl = rb_p2.add_run("Hinweis nach AHP Kap. 3.4: ")
    rb_p2_lbl.bold = True
    rb_p2_lbl.font.size = Pt(9)
    rb_p2_body = rb_p2.add_run(
        "Über Art und Höhe der Bewilligung entscheidet der Sozialausschuss "
        "der Stadt Würzburg nach Prüfung durch die Leitung Seniorenarbeit "
        "(Fachbereich IIS). Ein Rechtsanspruch auf Förderung besteht nicht."
    )
    rb_p2_body.font.size = Pt(9)

    # ── 12. FOOTER (Bescheid-ID + Doctree-Version) ─────────────────
    foot_p = doc.add_paragraph()
    foot_p.paragraph_format.space_before = Pt(8)
    foot_r = foot_p.add_run(
        f"Bescheid-ID: {bescheid_id} · "
        f"Erstellt auf Basis des Prüfprotokolls vom "
        f"{_format_date(ausgestellt_am)} · "
        f"Doctree-Version {doctree_version or '—'}"
    )
    foot_r.font.size = Pt(8)
    foot_r.font.color.rgb = RGBColor.from_string(_CI_HELLGRAU)

    # In-Memory-Bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
